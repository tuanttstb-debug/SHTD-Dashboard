"""
generate_docx.py — Tạo USER_MANUAL_FINAL.docx từ USER_MANUAL.md + ảnh HDSD
Run: python generate_docx.py
"""

import os, re, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE = Path(r"D:\Công việc\Vibecode\SHTD-Dashboard")
MD   = BASE / "USER_MANUAL.md"
HDSD = BASE / "HDSD"
OUT  = BASE / "USER_MANUAL_FINAL.docx"

# Screenshot catalog: marker in MD → (file, caption)
SCREENSHOT_MAP = {
    "HDSD/01_login.png":          (HDSD / "01_login.png",          "Hình 1 — Màn hình đăng nhập hệ thống"),
    "HDSD/02_dashboard.png":      (HDSD / "02_dashboard.png",      "Hình 2 — Tổng quan Dashboard"),
    "HDSD/03_create_task.png":    (HDSD / "03_create_task.png",    "Hình 3 — Form tạo Task mới"),
    "HDSD/03_create_case.png":    (HDSD / "03_create_case.png",    "Hình 4 — Form tạo Case mới"),
    "HDSD/04_edit_case.png":      (HDSD / "04_edit_case.png",      "Hình 5 — Form chỉnh sửa Case"),
    "HDSD/07_export.png":         (HDSD / "07_export.png",         "Hình 6 — Chức năng Xuất dữ liệu"),
    "HDSD/03_create_initiative.png": (HDSD / "03_create_initiative.png", "Hình 7 — Form tạo Initiative"),
    "HDSD/05_submit.png":         (HDSD / "05_submit.png",         "Hình 8 — Hàng đợi phê duyệt BLĐ"),
    "HDSD/06_approve.png":        (HDSD / "06_approve.png",        "Hình 9 — Modal phê duyệt BLĐ"),
    "HDSD/08_mobile_view.png":    (HDSD / "08_mobile_view.png",    "Hình 10 — Giao diện Mobile"),
}

# ── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # East Asian font
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)

def set_para_spacing(para, before=0, after=6, line=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type("page"))
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)

def docx_break_type(btype):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), btype)
    return br

def add_break_to_para(para, btype="page"):
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), btype)
    run._r.append(br)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), kwargs.get("val", "single"))
        border.set(qn("w:sz"), kwargs.get("sz", "4"))
        border.set(qn("w:color"), kwargs.get("color", "000000"))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_full_border_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)

def set_table_width(table, width_cm=16):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))  # 1cm = 567 twips
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

def shade_cell(cell, fill="E8F0FE"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def add_image_centered(doc, img_path, width_cm=14, caption=None):
    if not Path(img_path).exists():
        p = doc.add_paragraph(f"[Ảnh không tìm thấy: {img_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    set_para_spacing(para, before=6, after=3)
    if caption:
        cap_p = doc.add_paragraph(caption)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap_p.runs:
            set_font(run, size=10, italic=True, color=(89, 89, 89))
        if not cap_p.runs:
            r = cap_p.add_run(caption)
            set_font(r, size=10, italic=True, color=(89, 89, 89))
        set_para_spacing(cap_p, before=0, after=10)

def add_footer_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("Trang ")
        set_font(run, size=9, color=(128, 128, 128))
        # PAGE field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        r2 = OxmlElement("w:r")
        r2.append(fldChar1)
        r3 = OxmlElement("w:r")
        r3.append(instrText)
        r4 = OxmlElement("w:r")
        r4.append(fldChar2)
        para._p.append(r2)
        para._p.append(r3)
        para._p.append(r4)
        run2 = para.add_run(" / ")
        set_font(run2, size=9, color=(128, 128, 128))
        # NUMPAGES field
        fldChar5 = OxmlElement("w:fldChar")
        fldChar5.set(qn("w:fldCharType"), "begin")
        instrText2 = OxmlElement("w:instrText")
        instrText2.text = "NUMPAGES"
        fldChar6 = OxmlElement("w:fldChar")
        fldChar6.set(qn("w:fldCharType"), "end")
        r5 = OxmlElement("w:r")
        r5.append(fldChar5)
        r6 = OxmlElement("w:r")
        r6.append(instrText2)
        r7 = OxmlElement("w:r")
        r7.append(fldChar6)
        para._p.append(r5)
        para._p.append(r6)
        para._p.append(r7)

# ── Cover Page ─────────────────────────────────────────────────────────────

def build_cover(doc):
    # Spacer
    for _ in range(4):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0)

    # Bank name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KHỐI KHÁCH HÀNG DOANH NGHIỆP (KHDN)")
    set_font(r, size=12, color=(89, 89, 89))
    set_para_spacing(p, before=0, after=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BỘ PHẬN PHÁT TRIỂN KINH DOANH")
    set_font(r, size=11, color=(89, 89, 89))
    set_para_spacing(p, before=0, after=30)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HƯỚNG DẪN SỬ DỤNG")
    set_font(r, size=28, bold=True, color=(13, 71, 161))
    set_para_spacing(p, before=0, after=8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SHTD Dashboard")
    set_font(r, size=22, bold=True, color=(13, 71, 161))
    set_para_spacing(p, before=0, after=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Số Hóa Tín Dụng — Phiên bản v6.2")
    set_font(r, size=14, color=(25, 118, 210))
    set_para_spacing(p, before=0, after=40)

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("─" * 50)
    set_font(r, size=11, color=(189, 189, 189))
    set_para_spacing(p, before=0, after=20)

    # Info table
    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ("Phiên bản", "v6.2"),
        ("Ngày phát hành", "16/06/2026"),
        ("Đối tượng", "Nhân viên, Team Lead, Quản lý, BLĐ"),
        ("Bảo mật", "Nội bộ — Không phân phối bên ngoài"),
        ("Tác giả", "Bộ phận Chuyển đổi số KHDN"),
    ]
    for i, (label, value) in enumerate(rows_data):
        row = tbl.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        shade_cell(c0, "EEF2FF")
        p0 = c0.paragraphs[0]
        p0.clear()
        r0 = p0.add_run(label)
        set_font(r0, size=11, bold=True)
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1 = c1.paragraphs[0]
        p1.clear()
        r1 = p1.add_run(value)
        set_font(r1, size=11)
        set_cell_border(c0)
        set_cell_border(c1)
    add_full_border_table_borders(tbl)

    for _ in range(6):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tài liệu này được tạo tự động từ mã nguồn hệ thống SHTD Dashboard v6.2")
    set_font(r, size=9, italic=True, color=(158, 158, 158))
    set_para_spacing(p, before=0, after=0)

    add_break_to_para(doc.add_paragraph(), "page")

# ── TOC Placeholder ────────────────────────────────────────────────────────

def build_toc(doc):
    p = doc.add_heading("MỤC LỤC", level=1)
    for run in p.runs:
        set_font(run, size=16, bold=True)

    # TOC field instruction
    para = doc.add_paragraph()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r1 = OxmlElement("w:r")
    r1.append(fldChar1)
    r2 = OxmlElement("w:r")
    r2.append(instrText)
    r3 = OxmlElement("w:r")
    r3.append(fldChar2)
    # placeholder text
    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "[Nhấn Ctrl+A rồi F9 trong Microsoft Word để cập nhật Mục lục tự động]"
    r4.append(t)
    r5 = OxmlElement("w:r")
    r5.append(fldChar3)
    para._p.extend([r1, r2, r3, r4, r5])
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        set_font(run, size=10, italic=True, color=(128, 128, 128))

    add_break_to_para(doc.add_paragraph(), "page")

# ── MD Parser ──────────────────────────────────────────────────────────────

def parse_inline(text):
    """Return list of (text, bold, italic, code, link_url) tuples."""
    segments = []
    # Simplified: handle **bold**, *italic*, `code`, [text](url)
    pattern = re.compile(
        r'\*\*(.+?)\*\*'       # bold
        r'|\*(.+?)\*'          # italic
        r'|`([^`]+)`'          # inline code
        r'|\[([^\]]+)\]\(([^)]+)\)'  # link
        r'|(\[NEED HUMAN VALIDATION\])'  # special marker
    )
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            segments.append(("text", text[last:m.start()]))
        if m.group(1):
            segments.append(("bold", m.group(1)))
        elif m.group(2):
            segments.append(("italic", m.group(2)))
        elif m.group(3):
            segments.append(("code", m.group(3)))
        elif m.group(4):
            segments.append(("text", m.group(4)))
        elif m.group(6):
            segments.append(("marker", m.group(6)))
        last = m.end()
    if last < len(text):
        segments.append(("text", text[last:]))
    return segments

def apply_inline(para, text):
    for stype, val in parse_inline(text):
        if stype == "bold":
            r = para.add_run(val)
            set_font(r, bold=True)
        elif stype == "italic":
            r = para.add_run(val)
            set_font(r, italic=True)
        elif stype == "code":
            r = para.add_run(val)
            r.font.name = "Courier New"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(180, 0, 0)
        elif stype == "marker":
            r = para.add_run(val)
            set_font(r, bold=True, color=(255, 152, 0))
        else:
            r = para.add_run(val)
            set_font(r)

# ── Document Builder ───────────────────────────────────────────────────────

def build_body(doc, lines):
    i = 0
    in_table = False
    table_rows = []
    in_code = False
    code_lines = []
    in_list = False

    while i < len(lines):
        line = lines[i]
        raw  = line.rstrip()

        # ── Screenshot marker ──
        for marker, (img_path, caption) in SCREENSHOT_MAP.items():
            if marker in raw:
                add_image_centered(doc, img_path, width_cm=14, caption=caption)
                raw = ""
                break

        if not raw and not in_table and not in_code:
            i += 1
            continue

        # ── Code block ──
        if raw.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(50, 50, 50)
                p.paragraph_format.left_indent = Cm(1)
                # light gray shading
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F5F5F5")
                pPr.append(shd)
                set_para_spacing(p, before=4, after=4)
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # ── HR / separator ──
        if re.match(r'^[-─═*]{3,}\s*$', raw):
            p = doc.add_paragraph()
            set_para_spacing(p, before=4, after=4)
            i += 1
            continue

        # ── Headings ──
        h6 = re.match(r'^#{6}\s+(.*)', raw)
        h5 = re.match(r'^#{5}\s+(.*)', raw)
        h4 = re.match(r'^#{4}\s+(.*)', raw)
        h3 = re.match(r'^#{3}\s+(.*)', raw)
        h2 = re.match(r'^#{2}\s+(.*)', raw)
        h1 = re.match(r'^#{1}\s+(.*)', raw)

        if h1 and not h2:
            text = h1.group(1)
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                set_font(run, size=16, bold=True, color=(13, 71, 161))
            set_para_spacing(p, before=18, after=8)
            i += 1; continue

        if h2 and not h3:
            text = h2.group(1)
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                set_font(run, size=13, bold=True, color=(21, 101, 192))
            set_para_spacing(p, before=14, after=6)
            i += 1; continue

        if h3 and not h4:
            text = h3.group(1)
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                set_font(run, size=12, bold=True, color=(30, 136, 229))
            set_para_spacing(p, before=10, after=4)
            i += 1; continue

        if h4 and not h5:
            text = h4.group(1)
            p = doc.add_heading(text, level=4)
            for run in p.runs:
                set_font(run, size=11, bold=True, color=(66, 165, 245))
            set_para_spacing(p, before=8, after=3)
            i += 1; continue

        if h5 or h6:
            m = h5 or h6
            text = m.group(1)
            p = doc.add_paragraph()
            r = p.add_run(text)
            set_font(r, size=11, bold=True, color=(89, 89, 89))
            set_para_spacing(p, before=6, after=2)
            i += 1; continue

        # ── Tables (pipe-delimited) ──
        if raw.startswith("|"):
            table_rows.append(raw)
            i += 1
            # Peek ahead to collect full table
            while i < len(lines) and lines[i].rstrip().startswith("|"):
                table_rows.append(lines[i].rstrip())
                i += 1
            # Process table
            data_rows = [r for r in table_rows if not re.match(r'^\|[-:\s|]+\|$', r)]
            if data_rows:
                cells_per_row = [
                    [c.strip() for c in r.strip("|").split("|")]
                    for r in data_rows
                ]
                max_cols = max(len(r) for r in cells_per_row)
                tbl = doc.add_table(rows=len(cells_per_row), cols=max_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                add_full_border_table_borders(tbl)
                for ri, row_cells in enumerate(cells_per_row):
                    trow = tbl.rows[ri]
                    for ci, cell_text in enumerate(row_cells):
                        if ci >= max_cols:
                            break
                        cell = trow.cells[ci]
                        p = cell.paragraphs[0]
                        p.clear()
                        if ri == 0:
                            shade_cell(cell, "DBEAFE")
                            r = p.add_run(cell_text)
                            set_font(r, size=10, bold=True)
                        else:
                            apply_inline(p, cell_text)
                            for run in p.runs:
                                run.font.size = Pt(10)
                        set_para_spacing(p, before=2, after=2)
                doc.add_paragraph()
            table_rows = []
            continue

        # ── Blockquote ──
        if raw.startswith(">"):
            text = raw.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "FFF9C4")
            pPr.append(shd)
            apply_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.italic = True
            set_para_spacing(p, before=3, after=3)
            i += 1; continue

        # ── Unordered list ──
        ul = re.match(r'^(\s*)[-*+]\s+(.*)', raw)
        if ul:
            indent = len(ul.group(1)) // 2
            text = ul.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.8 + indent * 0.4)
            apply_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(11)
            set_para_spacing(p, before=1, after=1)
            i += 1; continue

        # ── Ordered list ──
        ol = re.match(r'^(\s*)\d+\.\s+(.*)', raw)
        if ol:
            indent = len(ol.group(1)) // 2
            text = ol.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.8 + indent * 0.4)
            apply_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(11)
            set_para_spacing(p, before=1, after=1)
            i += 1; continue

        # ── Normal paragraph ──
        if raw.strip():
            p = doc.add_paragraph()
            apply_inline(p, raw.strip())
            set_para_spacing(p, before=0, after=6)

        i += 1

# ── Validation Report ──────────────────────────────────────────────────────

def build_validation_report(doc):
    add_break_to_para(doc.add_paragraph(), "page")
    p = doc.add_heading("BÁO CÁO XÁC NHẬN TÀI LIỆU", level=1)
    for run in p.runs:
        set_font(run, size=16, bold=True, color=(13, 71, 161))

    items = [
        ("Tổng số trang", "Xem footer — cập nhật tự động"),
        ("Tổng số hình ảnh nhúng", f"{sum(1 for p, _ in SCREENSHOT_MAP.values() if p.exists())} / {len(SCREENSHOT_MAP)}"),
        ("Ảnh không tìm thấy", ", ".join(str(p.name) for p, _ in SCREENSHOT_MAP.values() if not p.exists()) or "Không có"),
        ("Nguồn nội dung", "USER_MANUAL.md (trích xuất từ mã nguồn)"),
        ("Phiên bản hệ thống", "SHTD Dashboard v6.2"),
        ("Ngày tạo tài liệu", "16/06/2026"),
        ("Công cụ tạo", "Python 3.14 + python-docx 1.2.0"),
        ("Cần xác nhận con người", "Các mục [NEED HUMAN VALIDATION] trong tài liệu"),
        ("Bộ phông chữ", "Calibri (body), Courier New (code)"),
        ("Cỡ chữ thân bài", "11pt"),
        ("Khoảng cách dòng", "1.15x"),
        ("Chiều rộng ảnh", "14cm (~70% trang A4)"),
    ]

    tbl = doc.add_table(rows=len(items), cols=2)
    add_full_border_table_borders(tbl)
    for i, (label, value) in enumerate(items):
        row = tbl.rows[i]
        shade_cell(row.cells[0], "EEF2FF")
        p0 = row.cells[0].paragraphs[0]
        p0.clear()
        r0 = p0.add_run(label)
        set_font(r0, size=10, bold=True)
        p1 = row.cells[1].paragraphs[0]
        p1.clear()
        r1 = p1.add_run(value)
        set_font(r1, size=10)
        set_para_spacing(p0, before=2, after=2)
        set_para_spacing(p1, before=2, after=2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "Ghi chú: Tài liệu này được tạo tự động. Nội dung phản ánh trạng thái hệ thống "
        "tại thời điểm tạo tài liệu. Các mục được đánh dấu [NEED HUMAN VALIDATION] "
        "cần được xác nhận bởi nhóm phát triển hoặc người dùng nghiệp vụ trước khi phân phối."
    )
    set_font(r, size=10, italic=True, color=(89, 89, 89))
    set_para_spacing(p, before=8, after=0)

# ── Main ───────────────────────────────────────────────────────────────────

def main():
    print("Đang tạo USER_MANUAL_FINAL.docx ...")

    doc = Document()

    # Page setup: A4
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    print("  [1/5] Trang bìa...")
    build_cover(doc)

    print("  [2/5] Mục lục...")
    build_toc(doc)

    print("  [3/5] Nội dung chính từ USER_MANUAL.md...")
    with open(MD, encoding="utf-8") as f:
        lines = f.readlines()

    # Skip the TOC section in MD (lines 15-36) since we have auto TOC
    # Keep everything else
    filtered = []
    skip_toc = False
    for line in lines:
        if "## MỤC LỤC" in line:
            skip_toc = True
            continue
        if skip_toc and line.strip().startswith("---"):
            skip_toc = False
            continue
        if skip_toc:
            continue
        filtered.append(line)

    build_body(doc, filtered)

    print("  [4/5] Báo cáo xác nhận...")
    build_validation_report(doc)

    print("  [5/5] Số trang footer...")
    add_footer_page_numbers(doc)

    doc.save(str(OUT))
    print(f"\n  ✅ Đã lưu: {OUT}")
    print(f"  Kích thước: {OUT.stat().st_size / 1024:.1f} KB")
    print("\nHoàn thành! Mở file trong Microsoft Word và nhấn Ctrl+A → F9 để cập nhật Mục lục.\n")

if __name__ == "__main__":
    main()
